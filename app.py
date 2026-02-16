from flask import Flask, render_template, request, jsonify, send_file
import speech_recognition as sr
import tempfile
import os
from indic_transliteration import sanscript
from indic_transliteration.sanscript import transliterate
from docx import Document
import io
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

# Allowed audio file extensions
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'm4a', 'flac', 'aac'}

# Unicode to Krutidev Mapping
UNICODE_TO_KRUTIDEV = {
    "अ": "v", "आ": "vk", "इ": "b", "ई": "bZ", "उ": "m", "ऊ": "Å",
    "ऋ": "_", "ए": ",", "ऐ": ",s", "ओ": "vks", "औ": "vkS",
    "क": "d", "ख": "[k", "ग": "x", "घ": "?k", "ङ": "³",
    "च": "p", "छ": "N", "ज": "t", "झ": ">", "ञ": "¥",
    "ट": "V", "ठ": "B", "ड": "M", "ढ": "<", "ण": ".k",
    "त": "r", "थ": "Fk", "द": "n", "ध": "/k", "न": "u",
    "प": "i", "फ": "Q", "ब": "c", "भ": "Hk", "म": "e",
    "य": ";", "र": "j", "ल": "y", "व": "o", "श": "'k",
    "ष": "\"k", "स": "l", "ह": "g",
    "ा": "k", "ि": "f", "ी": "h", "ु": "q", "ू": "w", "े": "s",
    "ै": "S", "ो": "ks", "ौ": "kS", "ं": "M+", "ः": "%", "ँ": "~",
    "्": "", " ": " ", "\n": "\n"
}


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def convert_unicode_to_krutidev(unicode_text):
    """Convert Unicode text to Krutidev format"""
    try:
        krutidev_text = ""
        i = 0
        while i < len(unicode_text):
            char = unicode_text[i]

            if i < len(unicode_text) - 1:
                compound = char + unicode_text[i + 1]
                if compound in UNICODE_TO_KRUTIDEV:
                    krutidev_text += UNICODE_TO_KRUTIDEV[compound]
                    i += 2
                    continue

            if char in UNICODE_TO_KRUTIDEV:
                krutidev_text += UNICODE_TO_KRUTIDEV[char]
            else:
                krutidev_text += char

            i += 1

        return krutidev_text

    except Exception as e:
        try:
            return transliterate(unicode_text, sanscript.DEVANAGARI, sanscript.ITRANS)
        except:
            return f"Conversion error: {str(e)}"


def recognize_speech(audio_file, language='mr-IN'):
    """Recognize speech from audio file"""
    r = sr.Recognizer()
    try:
        with sr.AudioFile(audio_file) as source:
            r.adjust_for_ambient_noise(source)
            audio_data = r.record(source)
            text = r.recognize_google(audio_data, language=language)
            return text

    except sr.UnknownValueError:
        return "Could not understand audio"
    except sr.RequestError as e:
        return f"Error with speech service: {e}"
    except Exception as e:
        return f"Error: {str(e)}"


@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')


@app.route('/transcribe', methods=['POST'])
def transcribe():
    """Handle audio transcription request"""
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        language = request.form.get('language', 'mr-IN')
        
        if audio_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if audio_file and allowed_file(audio_file.filename):
            # Save the uploaded file temporarily
            filename = secure_filename(audio_file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            audio_file.save(filepath)
            
            try:
                # Recognize speech
                unicode_text = recognize_speech(filepath, language)
                
                if unicode_text.startswith("Error") or unicode_text.startswith("Could not"):
                    return jsonify({'error': unicode_text}), 400
                
                # Convert to Krutidev
                krutidev_text = convert_unicode_to_krutidev(unicode_text)
                
                return jsonify({
                    'success': True,
                    'unicode_text': unicode_text,
                    'krutidev_text': krutidev_text
                })
                
            finally:
                # Clean up temporary file
                if os.path.exists(filepath):
                    os.unlink(filepath)
        
        return jsonify({'error': 'Invalid file type'}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/txt', methods=['POST'])
def download_txt():
    """Download text as TXT file"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        # Create a text file in memory
        text_io = io.BytesIO(text.encode('utf-8'))
        text_io.seek(0)
        
        return send_file(
            text_io,
            mimetype='text/plain',
            as_attachment=True,
            download_name='krutidev.txt'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/docx', methods=['POST'])
def download_docx():
    """Download text as DOCX file"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        # Create a Word document
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Kruti Dev 010'
        
        # Save to BytesIO object
        file_io = io.BytesIO()
        doc.save(file_io)
        file_io.seek(0)
        
        return send_file(
            file_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='krutidev.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    # Use PORT environment variable for Render
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)